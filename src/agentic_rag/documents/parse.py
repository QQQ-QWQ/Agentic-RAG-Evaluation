from pathlib import Path

from agentic_rag.documents.errors import UnsupportedFormatError
from agentic_rag.documents.models import ParsedDocument

_TEXT_SUFFIXES = frozenset({".txt", ".md", ".markdown"})


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def parse_path(path: str | Path) -> ParsedDocument:
    """
    将本地文档解析为纯文本。

    支持：.txt、.md、.pdf、.docx
    """
    p = Path(path).expanduser().resolve()
    if not p.is_file():
        raise FileNotFoundError(p)

    suffix = p.suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        text = _load_text(p)
    elif suffix == ".pdf":
        text = _load_pdf(p)
    elif suffix == ".docx":
        text = _load_docx(p)
    else:
        raise UnsupportedFormatError(suffix)

    return ParsedDocument(text=text, path=str(p), suffix=suffix)
